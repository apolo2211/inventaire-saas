# -*- coding: utf-8 -*-
from docx import Document
import pandas as pd
from reportlab.pdfgen import canvas
import io

# Générer un document Word pour un produit ou étiquette
def generate_word(product):
    doc = Document()
    doc.add_heading('Étiquette Produit', 0)
    doc.add_paragraph(f"Nom: {product['name']}")
    doc.add_paragraph(f"Description: {product['description']}")
    doc.add_paragraph(f"Quantité: {product['quantity']}")
    doc.add_paragraph(f"Code: {product['code_unique']}")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Générer un fichier Excel pour inventaire
def generate_excel(products):
    df = pd.DataFrame(products)
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)
    return buffer

# Générer un PDF simple pour rapport
def generate_pdf(products):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    c.drawString(100, 800, "Rapport Inventaire")
    y = 750
    for p in products:
        c.drawString(50, y, f"{p['name']} - {p['quantity']} unités - {p['value']} DZD")
        y -= 20
    c.save()
    buffer.seek(0)
    return buffer
